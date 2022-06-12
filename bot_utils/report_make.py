import os

from docx.enum.text import WD_BREAK
from docxtpl import DocxTemplate
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime
from docx.shared import Inches, Pt


def graphs(statistic,user):
    now = datetime.now()
    time = now.strftime("%m%d%Y%H%M%S")
    os.makedirs(f'stat/{time}')
    def pie(statistic,time):
        labels = 'Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс'
        sizes = statistic[1].split("")

        fig1, ax1 = plt.subplots()
        colors = plt.get_cmap('RdBu')(np.linspace(0.2, 1, len(sizes)))
        ax1.pie(sizes, labels=labels, autopct=lambda p: f'{p * sum(sizes) / 100 :.0f} score', colors=colors,
                startangle=90)
        ax1.axis('equal')

        plt.savefig(f'stat/{time}/pie.png')
    def bars(statistic,time):
        fig, ax = plt.subplots()

        day = ('Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс')
        y_pos = np.arange(len(day))
        performance = statistic[2].split('')

        rects1 = ax.barh(y_pos, performance, align='center', )
        ax.bar_label(rects1, padding=3)
        ax.set_yticks(y_pos, labels=day)
        ax.invert_yaxis()  # labels read top-to-bottom
        ax.set_xlabel('Оценка')
        ax.set_title('Твое настроение за неделю')

        plt.savefig(f'stat/{time}/bars.png')


    def mounts(statistic,time):
        names = ('Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс')
        values =statistic[3].split("")

        fig, axs = plt.subplots()
        axs.plot(names, values)
        fig.suptitle('Categorical Plotting')
        plt.savefig(f'stat/{time}/mounts.png')
    def document_save(user,time):
        tmplt=DocxTemplate("report.docx")
        stat=statistic[0].split("@")
        context = {"subject": stat[0],
                   "thema_now": stat[1],
                   "learnA": stat[2],
                   "rememberB": stat[3],
                   "task": stat[4],
                   }
        tmplt.render(context)
        # "Карта эмоционального состояния Уровень мотивации Регулярность занятий"

        para = tmplt.add_paragraph().add_run(
            'Регулярность занятий')
        para.font.size = Pt(18)
        para.name = 'PT Sans Narrow'
        para.bold = True
        tmplt.add_picture('stat/06122022203107/bars.png', width=Inches(3.5))

        para = tmplt.add_paragraph().add_run(
            'Карта эмоционального состояния')
        para.font.size = Pt(18)
        para.name = 'PT Sans Narrow'
        para.bold = True
        tmplt.add_picture('stat/06122022203107/mounts.png', width=Inches(3.5))
        para = tmplt.add_paragraph().add_run(
            'Уровень мотивации')
        para.font.size = Pt(18)
        para.name = 'PT Sans Narrow'
        para.bold = True
        tmplt.add_picture('stat/06122022203107/pie.png', width=Inches(3.5))
        tmplt.save("generated_doc.docx")

    pie(statistic,time)
    bars(statistic,time)
    mounts(statistic,time)
    document_save(user, time)
    os.remove(f'stat/{time}')